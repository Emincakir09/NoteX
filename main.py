# ==========================================
# 🚨 KRİTİK YAMA (BU KISIM EN ÜSTTE KALMALI)
# ==========================================
import sys
from unittest.mock import MagicMock

# 1. requests_toolbelt Yaması
try:
    import requests_toolbelt.adapters.appengine
except (ImportError, ModuleNotFoundError):
    sys.modules["requests_toolbelt.adapters.appengine"] = MagicMock()

# 2. urllib3 Yaması
m = MagicMock()
m.is_appengine_sandbox = lambda: False
sys.modules["urllib3.contrib.appengine"] = m
# ==========================================

import streamlit as st
import firebase_admin
from firebase_admin import credentials, firestore, storage
import pyrebase
import os
from dotenv import load_dotenv
import PyPDF2
import json
import time
import streamlit.components.v1 as components 
from pyvis.network import Network 

# --- YAPAY ZEKA KÜTÜPHANELERİ ---
import google.generativeai as genai
from langchain_core.messages import HumanMessage, AIMessage

# --- RAG SERVİSİ (YENİ) ---
try:
    from lib.rag_service import RAGService
except ImportError as e:
    st.error(f"RAG kütüphaneleri yüklenemedi: {e}")
    st.stop()

# 1. Ayarları Yükle
load_dotenv()
api_key = os.getenv("GOOGLE_API_KEY")

if not api_key:
    st.warning("⚠️ Google API Key bulunamadı! .env dosyanızı kontrol edin.")
else:
    genai.configure(api_key=api_key)

st.set_page_config(page_title="Akademik Agent (RAG Pro)", page_icon="🌙", layout="wide", initial_sidebar_state="auto")

# --- CSS (KOYU TEMA & DÜZELTMELER) ---
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&display=swap');

    /* Genel Arkaplan ve Metin Renkleri */
    body {
        color: #1F2937;
        background-color: #FFFFFF; /* Tamamen beyaz arka plan daha ferah durur */
        font-family: 'Inter', sans-serif;
    }
    
    /* Streamlit varsayılan arkaplanlarını ezmek için */
    .stApp {
        background-color: #FFFFFF !important;
    }

    /* Sohbet Balonları */
    .stChatMessage {
        padding: 1.2rem 1.5rem;
        border-radius: 16px;
        margin-bottom: 20px;
        border: 1px solid #F3F4F6;
        box-shadow: 0 2px 4px rgba(0,0,0,0.02);
        font-family: 'Inter', sans-serif;
    }
    
    /* Kullanıcı (Sağ) */
    [data-testid="stChatMessage"]:nth-child(odd) {
        background-color: #F8FAFC !important; /* Açık grimsi mavi */
        border: 1px solid #E2E8F0 !important;
        border-left: 4px solid #3B82F6 !important;
    }

    /* Asistan (Sol) */
    [data-testid="stChatMessage"]:nth-child(even) {
        background-color: #FFFFFF !important; /* Temiz Beyaz */
        border: 1px solid #E5E7EB !important;
        border-right: 4px solid #10B981 !important; /* Zümrüt yeşili vurgu */
    }

    /* Metin Renkleri (Okunabilirlik İçin) */
    [data-testid="stChatMessage"] p, 
    [data-testid="stChatMessage"] div,
    [data-testid="stChatMessage"] span,
    [data-testid="stChatMessage"] li {
        color: #374151 !important;
        font-family: 'Inter', sans-serif !important;
        font-size: 15px;
        line-height: 1.6;
    }
    
    code {
        color: #EF4444 !important; /* Kırmızımsı kod vurgusu */
        background-color: #F3F4F6 !important;
        border-radius: 6px;
        padding: 0.2em 0.4em;
        font-family: monospace;
        font-size: 0.9em;
    }

    /* Butonlar */
    .stButton>button {
        width: 100%; 
        border-radius: 10px; 
        height: 3.2em; 
        font-weight: 600;
        background-color: #FFFFFF;
        color: #374151;
        border: 1px solid #D1D5DB;
        transition: all 0.2s ease-in-out;
        box-shadow: 0 1px 2px rgba(0,0,0,0.05);
        font-family: 'Inter', sans-serif;
    }
    .stButton>button:hover {
        border-color: #3B82F6;
        color: #3B82F6;
        background-color: #F8FAFC;
        transform: translateY(-1px);
        box-shadow: 0 4px 6px rgba(59, 130, 246, 0.1);
    }
    
    /* Primary Buton */
    .stButton>button[data-baseweb="button"] {
        border-color: #E5E7EB;
    }

    /* --- NOTEBOOK KARTI TASARIMI (YENİ) --- */
    .notebook-card {
        background-color: #FFFFFF;
        padding: 24px;
        border-radius: 12px;
        font-family: 'Inter', sans-serif;
        font-size: 15px;
        line-height: 1.7;
        color: #374151;
        box-shadow: 0 4px 6px -1px rgba(0,0,0,0.05), 0 2px 4px -1px rgba(0,0,0,0.03);
        margin-top: 20px;
        width: 100%; 
        border: 1px solid #E5E7EB;
        border-left: 6px solid #3B82F6;
    }

    .notebook-header {
        font-weight: 700;
        color: #1F2937;
        margin-bottom: 16px;
        font-size: 1.25rem;
        border-bottom: 2px solid #F3F4F6;
        padding-bottom: 12px;
        display: flex;
        align-items: center;
        gap: 12px;
    }

    .notebook-content {
        white-space: pre-wrap; 
        color: #4B5563;
    }

    /* Sidebar Başlıkları */
    .st-emotion-cache-vk329t, .st-emotion-cache-1rf8qj6 h1, .st-emotion-cache-1rf8qj6 h2, .st-emotion-cache-1rf8qj6 h3 { 
        color: #111827 !important; 
        font-weight: 700 !important;
        font-family: 'Inter', sans-serif;
    }

    /* ── NotebookLM-Style Nav Cards (Radio) ── */
    .nb-nav-radio div[data-testid="stRadio"] > label { display: none; }
    .nb-nav-radio div[data-testid="stRadio"] > div {
        display: flex !important;
        flex-direction: row !important;
        gap: 12px !important;
        width: 100%;
        padding-bottom: 10px;
    }
    .nb-nav-radio div[data-testid="stRadio"] > div > label {
        flex: 1 !important;
        height: 86px !important;
        min-height: 86px !important;
        max-height: 86px !important;
        background: #FFFFFF !important;
        border: 1px solid #E5E7EB !important;
        border-radius: 16px !important;
        display: flex !important;
        align-items: center !important;
        justify-content: center !important;
        cursor: pointer !important;
        font-weight: 600 !important;
        font-size: 0.9rem !important;
        color: #4B5563 !important;
        transition: all 0.2s ease !important;
        padding: 0 12px !important;
        margin: 0 !important;
        text-align: center !important;
        box-shadow: 0 1px 2px rgba(0,0,0,0.02) !important;
    }
    /* Radio circle gizle */
    .nb-nav-radio div[data-testid="stRadio"] > div > label > div:first-child {
        display: none !important;
    }
    .nb-nav-radio div[data-testid="stRadio"] > div > label:hover {
        border-color: #9CA3AF !important;
        background: #F9FAFB !important;
        transform: translateY(-2px) !important;
        box-shadow: 0 4px 6px rgba(0,0,0,0.04) !important;
    }
    .nb-nav-radio div[data-testid="stRadio"] > div > label:has(input:checked) {
        background: #EFF6FF !important;
        border-color: #3B82F6 !important;
        border-width: 2px !important;
        color: #1E3A8A !important;
        box-shadow: 0 4px 12px rgba(59, 130, 246, 0.15) !important;
    }

    /* Profil Kutusu Stilleri */
    .profile-box {
        background-color: #FFFFFF; 
        border-radius: 12px;
        padding: 20px;
        margin-top: 24px;
        border: 1px solid #E5E7EB;
        box-shadow: 0 2px 4px rgba(0,0,0,0.02);
        color: #374151;
        font-family: 'Inter', sans-serif;
    }

    .profile-box h4 {
        color: #111827; 
        margin-bottom: 12px;
        font-size: 1.15em;
        font-weight: 700;
    }

    .profile-box p {
        font-size: 0.95em;
        margin-bottom: 8px;
        color: #4B5563;
    }

    .profile-box .stButton>button {
        background-color: #FEE2E2; 
        border-color: #FECACA;
        color: #DC2626;
        margin-top: 10px;
        font-weight: 700;
    }

    .profile-box .stButton>button:hover {
        background-color: #DC2626;
        border-color: #DC2626;
        color: white;
    }

    /* Streamlit ana içeriği için padding */
    .main .block-container {
        padding-top: 3rem;
        padding-right: 3rem;
        padding-left: 3rem;
        padding-bottom: 3rem;
        max-width: 1000px;
    }

    /* Sidebar tasarımı */
    [data-testid="stSidebar"] {
        background-color: #F8FAFC !important;
        border-right: 1px solid #E2E8F0;
    }

    /* Input alanları */
    .stTextInput>div>div>input, .stTextArea>div>div>textarea {
        background-color: #FFFFFF;
        color: #1F2937;
        border: 1px solid #D1D5DB;
        border-radius: 10px;
        padding: 12px;
        font-size: 15px;
        box-shadow: 0 1px 2px rgba(0,0,0,0.02) inset;
    }
    .stTextInput>div>div>input:focus, .stTextArea>div>div>textarea:focus {
        border-color: #3B82F6;
        box-shadow: 0 0 0 2px rgba(59, 130, 246, 0.2);
    }

    /* Selectbox ve Multiselect */
    .stSelectbox>div>div, .stMultiSelect>div>div {
        background-color: #FFFFFF;
        color: #1F2937;
        border: 1px solid #D1D5DB;
        border-radius: 10px;
    }
    .stSelectbox>div>div:focus, .stMultiSelect>div>div:focus {
        border-color: #3B82F6;
        box-shadow: 0 0 0 2px rgba(59, 130, 246, 0.2);
    }

    /* Checkbox */
    .stCheckbox>label {
        color: #374151;
        font-weight: 500;
    }

    /* Başlıklar */
    h1, h2, h3, h4, h5, h6 {
        color: #111827 !important; 
        font-family: 'Inter', sans-serif !important;
        font-weight: 800 !important;
        letter-spacing: -0.02em;
    }

    /* Uyarı ve Hata Mesajları */
    .stAlert {
        border-radius: 10px;
        font-family: 'Inter', sans-serif;
    }

    /* ── BİLGİ KARTI (FLIP CARD) ── */
    .fc-wrapper {
        perspective: 1200px;
        width: 100%;
        margin: 0 auto;
        cursor: pointer;
    }
    .fc-inner {
        position: relative;
        width: 100%;
        min-height: 320px;
        transform-style: preserve-3d;
        transition: transform 0.6s cubic-bezier(0.4, 0.2, 0.2, 1);
    }
    .fc-inner.flipped { transform: rotateY(180deg); }
    .fc-front, .fc-back {
        position: absolute;
        width: 100%;
        min-height: 320px;
        backface-visibility: hidden;
        border-radius: 20px;
        display: flex;
        flex-direction: column;
        align-items: center;
        justify-content: center;
        padding: 40px;
        box-sizing: border-box;
        text-align: center;
        font-family: 'Inter', sans-serif;
    }
    .fc-front {
        background: #FFFFFF;
        border: 1px solid #E5E7EB;
        box-shadow: 0 10px 25px -5px rgba(0,0,0,0.05), 0 8px 10px -6px rgba(0,0,0,0.01);
    }
    .fc-back {
        background: linear-gradient(135deg, #1E3A8A 0%, #3B82F6 100%);
        border: none;
        box-shadow: 0 10px 25px -5px rgba(59, 130, 246, 0.3);
        transform: rotateY(180deg);
    }
    .fc-category {
        font-size: 0.75rem;
        font-weight: 700;
        letter-spacing: 0.1em;
        text-transform: uppercase;
        background: #EEF2FF; /* Indigo 50 */
        color: #4F46E5; /* Indigo 600 */
        border: 1px solid #C7D2FE; /* Indigo 200 */
        border-radius: 999px;
        padding: 4px 14px;
        margin-bottom: 24px;
    }
    .fc-back .fc-category {
        background: rgba(255,255,255,0.2);
        color: #FFFFFF;
        border-color: rgba(255,255,255,0.3);
    }
    .fc-term {
        font-size: 2rem;
        font-weight: 800;
        color: #111827;
        line-height: 1.3;
        letter-spacing: -0.02em;
    }
    .fc-hint {
        font-size: 0.85rem;
        color: #6B7280;
        margin-top: 30px;
        font-weight: 500;
    }
    .fc-back .fc-hint { color: #DBEAFE; } /* Blue 100 */
    .fc-definition {
        font-size: 1.15rem;
        color: #FFFFFF;
        line-height: 1.8;
        font-weight: 500;
    }
    .fc-counter {
        text-align: center;
        font-size: 0.9rem;
        color: #6B7280;
        margin-bottom: 12px;
        font-weight: 600;
        font-family: 'Inter', sans-serif;
    }
    .fc-progress {
        width: 100%; height: 6px;
        background: #E5E7EB;
        border-radius: 3px;
        margin-bottom: 24px;
        overflow: hidden;
    }
    .fc-progress-bar {
        height: 6px;
        background: #3B82F6;
        border-radius: 3px;
        transition: width 0.5s cubic-bezier(0.4, 0, 0.2, 1);
    }
    .card-list-row {
        display: flex;
        gap: 16px;
        align-items: flex-start;
        padding: 16px 0;
        border-bottom: 1px solid #F3F4F6;
    }
    .card-list-num {
        min-width: 36px;
        font-weight: 800;
        color: #3B82F6;
        font-size: 1.1rem;
        font-family: 'Inter', sans-serif;
    }
    .card-list-term { 
        font-weight: 700; 
        color: #1F2937; 
        font-family: 'Inter', sans-serif;
        font-size: 1.05rem;
    }
    .card-list-def  { 
        font-size: 0.95rem; 
        color: #4B5563; 
        margin-top: 4px; 
        line-height: 1.6;
    }
    .card-list-cat  {
        font-size: 0.7rem; 
        text-transform: uppercase; 
        letter-spacing: 0.05em;
        color: #4F46E5; 
        background: #EEF2FF;
        border-radius: 999px; 
        padding: 2px 10px;
        display: inline-block; 
        margin-top: 8px;
        font-weight: 600;
    }
</style>
""", unsafe_allow_html=True)

# ------------------------------------------------------------------
# 🔑 FIREBASE AYARLARI
# ------------------------------------------------------------------
firebaseConfig = {
  "apiKey": "AIzaSyARecm2N4o4sFsKX4_cso4ASZE0aVTlIY4",
  "authDomain": "akademikagent.firebaseapp.com",
  "projectId": "akademikagent",
  "storageBucket": "akademikagent.firebasestorage.app",
  "messagingSenderId": "1078772904266",
  "appId": "1:1078772904266:web:74e0e86b4418f0c287d1b3",
  "databaseURL": "" 
}

# Başlatmalar
try:
    firebase_auth = pyrebase.initialize_app(firebaseConfig)
    auth = firebase_auth.auth()
except Exception as e:
    st.error(f"Pyrebase Hatası: {e}")

if not firebase_admin._apps:
    try:
        cred = credentials.Certificate("serviceAccountKey.json")
        firebase_admin.initialize_app(cred, {
            'storageBucket': 'akademikagent.firebasestorage.app'
        })
    except Exception as e:
        pass

try:
    db = firestore.client()
    bucket = storage.bucket()
except:
    db = None
    bucket = None

# ════════════════════════════════════════════
# 🌐 ÇEVİRİ SİSTEMİ
# ════════════════════════════════════════════
TRANSLATIONS = {
    # ── GENEL / SIDEBAR ──
    "app_title":        {"TR": "🎓 AkademikAgent",                   "EN": "🎓 AcademicAgent"},
    "app_subtitle":     {"TR": "AI Destekli Öğrenme Platformu",       "EN": "AI-Powered Learning Platform"},
    "role_student":     {"TR": "Öğrenci",                             "EN": "Student"},
    "logout":           {"TR": "🚪 Çıkış Yap",                       "EN": "🚪 Log Out"},
    "chat_history":     {"TR": "🗂️ Sohbet Geçmişi",                  "EN": "🗂️ Chat History"},
    "new_chat":         {"TR": "➕ Yeni Sohbet",                     "EN": "➕ New Chat"},
    "no_chat_yet":      {"TR": "Henüz sohbet yok.",                   "EN": "No chats yet."},
    "documents":        {"TR": "📚 Belgeler",                         "EN": "📚 Documents"},
    "select_all":       {"TR": "Tümünü Seç",                          "EN": "Select All"},
    "clear":            {"TR": "Temizle",                             "EN": "Clear"},
    "no_doc_yet":       {"TR": "Henüz belge yüklenmedi.",             "EN": "No documents uploaded yet."},
    "upload_doc":       {"TR": "⬆️ Yeni Belge Yükle",               "EN": "⬆️ Upload Document"},
    "source_type":      {"TR": "Kaynak Tipi",                         "EN": "Source Type"},
    "pdf_label":        {"TR": "📄 PDF",                              "EN": "📄 PDF"},
    "audio_label":      {"TR": "🎙️ Ses",                             "EN": "🎙️ Audio"},
    "choose_pdf":       {"TR": "PDF Seç",                             "EN": "Choose PDF"},
    "upload_learn":     {"TR": "📥 Yükle & Öğren",                   "EN": "📥 Upload & Learn"},
    "choose_audio":     {"TR": "Ses Dosyası",                         "EN": "Audio File"},
    "process_learn":    {"TR": "🔊 İşle & Öğren",                   "EN": "🔊 Process & Learn"},
    "processing":       {"TR": "İşleniyor…",                          "EN": "Processing…"},
    "reading":          {"TR": "Okunuyor ve öğreniliyor…",            "EN": "Reading and learning…"},
    # ── HEADER ──
    "header_title":     {"TR": "🎓 Akademik Asistan",                 "EN": "🎓 Academic Assistant"},
    "header_sub":       {"TR": "Gemini Flash · FAISS RAG",            "EN": "Gemini Flash · FAISS RAG"},
    "badge_active":     {"TR": "{n} belge aktif",                     "EN": "{n} documents active"},
    "badge_none":       {"TR": "⚠️ Belge seçili değil",              "EN": "⚠️ No document selected"},
    # ── NAV CARDS ──
    "nav_chat_title":   {"TR": "Akıllı Sohbet",                      "EN": "Smart Chat"},
    "nav_chat_desc":    {"TR": "RAG destekli soru-cevap",             "EN": "RAG-powered Q&A"},
    "nav_quiz_title":   {"TR": "Sınav Hazırla",                      "EN": "Take a Quiz"},
    "nav_quiz_desc":    {"TR": "Belgeden otomatik quiz",              "EN": "Auto quiz from docs"},
    "nav_map_title":    {"TR": "Kavram Haritası",                     "EN": "Concept Map"},
    "nav_map_desc":     {"TR": "İlişki ağı & analiz",                "EN": "Relation graph & analysis"},
    "nav_analytics_title": {"TR": "Quiz Analizi",                     "EN": "Quiz Analytics"},
    "nav_analytics_desc":  {"TR": "Skor geçmişi & grafik",           "EN": "Score history & chart"},
    # ── SOHBET ──
    "up_doc_first":     {"TR": "👈 Sol menüden bir PDF veya ses dosyası yükleyin.", "EN": "👈 Upload a PDF or audio file from the left menu."},
    "sel_doc_first":    {"TR": "📚 Sol menüden en az bir belge seçin.", "EN": "📚 Select at least one document from the left menu."},
    "chat_placeholder": {"TR": "Ders notlarınızdan sorunuzu sorun…",  "EN": "Ask a question from your study notes…"},
    "researching":      {"TR": "Araştırılıyor…",                      "EN": "Researching…"},
    # ── SINAV ──
    "quiz_title":       {"TR": "📝 Sınav Hazırla",                   "EN": "📝 Prepare a Quiz"},
    "quiz_sel_doc":     {"TR": "Sol menüden sınav için belge seçin.", "EN": "Select documents for the quiz from the left menu."},
    "quiz_create_btn":  {"TR": "🚀 Yeni Sınav Oluştur  ({n} Belge)", "EN": "🚀 Generate New Quiz  ({n} Docs)"},
    "quiz_spinner":     {"TR": "Tüm belgeler analiz ediliyor, sorular hazırlanıyor…", "EN": "Analyzing all documents, preparing questions…"},
    "quiz_progress":    {"TR": "İlerleme: {a}/{t} soru cevaplandı",  "EN": "Progress: {a}/{t} questions answered"},
    "quiz_answer_btn":  {"TR": "Cevapla ({i})",                      "EN": "Answer ({i})"},
    "quiz_choose_opt":  {"TR": "Bir seçenek işaretleyin.",            "EN": "Please select an option."},
    "quiz_correct":     {"TR": "✅ Doğru!",                           "EN": "✅ Correct!"},
    "quiz_wrong":       {"TR": "❌ Yanlış. Doğru: {ans}",            "EN": "❌ Wrong. Correct: {ans}"},
    "quiz_finish":      {"TR": "🏁 Sınavı Tamamla ve Puanı Kaydet", "EN": "🏁 Finish Quiz & Save Score"},
    "quiz_warn_unanswered": {"TR": "Henüz {n} soru cevaplanmadı. Devam etmek istiyor musunuz?", "EN": "{n} questions left unanswered. Do you want to proceed?"},
    "quiz_force_submit":{"TR": "Evet, yine de bitir",                 "EN": "Yes, finish anyway"},
    "quiz_success_rate":{"TR": "Başarı Oranı: %{pct}",               "EN": "Success Rate: {pct}%"},
    "quiz_answer_key":  {"TR": "📋 Cevap Anahtarı",                  "EN": "📋 Answer Key"},
    "quiz_correct_ans": {"TR": "← Doğru Cevap",                      "EN": "← Correct Answer"},
    "quiz_retry":       {"TR": "🔄 Tekrar Çöz",                      "EN": "🔄 Retry Quiz"},
    # ── HARİTA ──
    "map_title":        {"TR": "🕸️ Kavram Haritası",                 "EN": "🕸️ Concept Map"},
    "map_sel_doc":      {"TR": "Sol menüden harita için belge seçin.","EN": "Select documents for the map from the left menu."},
    "map_draw":         {"TR": "🗺️ Haritayı Çiz",                   "EN": "🗺️ Draw Map"},
    "map_drawing":      {"TR": "Kavramlar çıkarılıyor…",             "EN": "Extracting concepts…"},
    "map_analysis_title": {"TR": "🔗 Kavram Bağlantı Analizi",       "EN": "🔗 Concept Link Analysis"},
    "map_analysis_desc":  {"TR": "Haritadaki iki kavramı seçin, yapay zeka aralarındaki ilişkiyi açıklasın.", "EN": "Select two concepts from the map and let AI explain the relationship."},
    "node1_label":      {"TR": "1. Kavram",                           "EN": "1st Concept"},
    "node2_label":      {"TR": "2. Kavram",                           "EN": "2nd Concept"},
    "analyze_btn":      {"TR": "🔍 Analiz",                          "EN": "🔍 Analyze"},
    "analyzing":        {"TR": "İnceleniyor…",                        "EN": "Analyzing…"},
    # ── ANALİTİK ──
    "analytics_title":  {"TR": "📊 Quiz Geçmişi & Analiz",           "EN": "📊 Quiz History & Analytics"},
    "analytics_no_data":{"TR": "Henüz tamamlanmış sınav yok. Sınav yapıp 'Sınavı Tamamla' butonuna basın.", "EN": "No completed quizzes yet. Take a quiz and press 'Finish Quiz'."},
    # ── BİLGİ KARTLARI ──
    "nav_cards_title":  {"TR": "Bilgi Kartları",                              "EN": "Knowledge Cards"},
    "nav_cards_desc":   {"TR": "Kavram kartları oluştur",                     "EN": "Generate concept cards"},
    "cards_title":      {"TR": "🃏 Bilgi Kartları",                           "EN": "🃏 Knowledge Cards"},
    "cards_sel_doc":    {"TR": "Sol menüden belge seçin.",                    "EN": "Select documents from the left menu."},
    "cards_create_btn": {"TR": "🧠 Kartları Oluştur  ({n} Belge)",           "EN": "🧠 Generate Cards  ({n} Docs)"},
    "cards_spinner":    {"TR": "Belgeler analiz ediliyor, kartlar hazırlanıyor…", "EN": "Analyzing documents, generating cards…"},
    "cards_flip":       {"TR": "↩ Kartı Çevir",                              "EN": "↩ Flip Card"},
    "cards_prev":       {"TR": "⬅ Önceki",                                   "EN": "⬅ Prev"},
    "cards_next":       {"TR": "Sonraki ➡",                                  "EN": "Next ➡"},
    "cards_of":         {"TR": "{cur} / {total} Kart",                       "EN": "Card {cur} / {total}"},
    "cards_export_pdf": {"TR": "📄 PDF İndir",                               "EN": "📄 Download PDF"},
    "cards_export_doc": {"TR": "📝 Word İndir",                              "EN": "📝 Download Word"},
    "cards_no_cards":   {"TR": "⚠️ Kart oluşturulamadı. Metni kontrol edin.","EN": "⚠️ Could not generate cards. Check the document."},
    "cards_all":        {"TR": "📋 Tüm Kartlar",                             "EN": "📋 All Cards"},
    "cards_regenerate": {"TR": "🔄 Yeni Kartlar Üret",                      "EN": "🔄 Regenerate Cards"},
    "metric_total":     {"TR": "📋 Toplam Sınav",                    "EN": "📋 Total Quizzes"},
    "metric_avg":       {"TR": "📈 Ortalama Başarı",                 "EN": "📈 Average Score"},
    "metric_best":      {"TR": "🏆 En Yüksek Skor",                  "EN": "🏆 Best Score"},
    "metric_last":      {"TR": "🕐 Son Sınav",                       "EN": "🕐 Last Quiz"},
    "trend_title":      {"TR": "📈 Başarı Trendi",                   "EN": "📈 Performance Trend"},
    "chart_col":        {"TR": "Başarı (%)",                          "EN": "Score (%)"},
    "quiz_label":       {"TR": "Sınav {i}",                          "EN": "Quiz {i}"},
    "band_excellent":   {"TR": "🏆 %80+ Mükemmel",                   "EN": "🏆 80%+ Excellent"},
    "band_improving":   {"TR": "📈 %50–79 Gelişiyor",                "EN": "📈 50–79% Improving"},
    "band_retry":       {"TR": "📉 %0–49 Tekrar Gerekli",            "EN": "📉 0–49% Needs Work"},
    "detail_title":     {"TR": "📋 Sınav Detayları",                 "EN": "📋 Quiz Details"},
    "detail_correct":   {"TR": "{s}/{t} doğru",                      "EN": "{s}/{t} correct"},
    # ── ARAŞTIRMA DEFTERLERİ ──
    "researches":        {"TR": "📂 Araştırmalarım",                "EN": "📂 My Research"},
    "new_research":      {"TR": "+ Yeni Araştırma",               "EN": "+ New Research"},
    "no_research_yet":   {"TR": "Henüz araştırma yok.",           "EN": "No research yet."},
    "research_sources":  {"TR": "📎 Araştırma Kaynakları",        "EN": "📎 Research Sources"},
    "untitled_research": {"TR": "Yeni Araştırma",                 "EN": "New Research"},
    "research_title_ph": {"TR": "Araştırma başlığı...",           "EN": "Research title..."},
    "del_research":      {"TR": "Araştırmayı sil",               "EN": "Delete research"},
    "no_source_yet":     {"TR": "Kaynak eklenmemiş.",             "EN": "No sources added yet."},
    # ── LOGIN ──
    "login_title":      {"TR": "Giriş",                               "EN": "Login"},
    "login_tab":        {"TR": "Giriş",                               "EN": "Login"},
    "register_tab":     {"TR": "Kayıt",                               "EN": "Register"},
    "email_label":      {"TR": "E-posta",                             "EN": "Email"},
    "pass_label":       {"TR": "Şifre",                               "EN": "Password"},
    "login_btn":        {"TR": "Giriş",                               "EN": "Log In"},
    "register_btn":     {"TR": "Kayıt Ol",                            "EN": "Register"},
    # ── POMODORO ──
    "pomodoro_title":   {"TR": "Odaklanma Modu",                      "EN": "Focus Mode"},
    "pomodoro_start":   {"TR": "Başla",                               "EN": "Start"},
    "pomodoro_reset":   {"TR": "Sıfırla",                             "EN": "Reset"},
    "pomodoro_break":   {"TR": "Mola zamanı! Süre doldu.",            "EN": "Time's up! Take a break."},
}

def t(key: str, **kwargs) -> str:
    """Aktif dile göre çeviriyi döndürür. Eksik anahtarda key'i gösterir."""
    lang = st.session_state.get("lang", "TR")
    text = TRANSLATIONS.get(key, {}).get(lang, f"[{key}]")
    for k, v in kwargs.items():
        text = text.replace(f"{{{k}}}", str(v))
    return text

# ════════════════════════════════════════════
# --- SESSION STATE ---
if "user" not in st.session_state: st.session_state.user = None
if "chat_history" not in st.session_state: st.session_state.chat_history = []
if "quiz_data" not in st.session_state: st.session_state.quiz_data = None
if "network_html" not in st.session_state: st.session_state.network_html = None
if "network_data" not in st.session_state: st.session_state.network_data = []
if "rag_ready" not in st.session_state: st.session_state.rag_ready = False
if "selected_rag_docs" not in st.session_state: st.session_state.selected_rag_docs = []
if "notebook_explanation" not in st.session_state: st.session_state.notebook_explanation = None
if "current_chat_id" not in st.session_state: st.session_state.current_chat_id = None
if "active_page" not in st.session_state: st.session_state.active_page = "chat"  # "chat" | "quiz" | "map" | "analytics"
if "quiz_answers" not in st.session_state: st.session_state.quiz_answers = {}  # {soru_idx: True/False}
if "quiz_submitted" not in st.session_state: st.session_state.quiz_submitted = False
if "lang" not in st.session_state: st.session_state.lang = "TR"  # "TR" | "EN"
if "flashcards" not in st.session_state: st.session_state.flashcards = []
if "card_index" not in st.session_state: st.session_state.card_index = 0
if "card_flipped" not in st.session_state: st.session_state.card_flipped = False
if "current_research_sources" not in st.session_state: st.session_state.current_research_sources = []


# --- RAG BAŞLATMA ---
def init_rag_for_user(email):
    user_db_path = f"./chroma_db/{email.replace('@', '_').replace('.', '_')}"
    return RAGService(persistence_path=user_db_path)

if st.session_state.user and "rag_service" not in st.session_state:
    st.session_state.rag_service = init_rag_for_user(st.session_state.user['email'])
    st.session_state.rag_ready = True

# --- FIREBASE FONKSİYONLARI ---
def upload_to_firebase(file_bytes, file_name, file_type, user_email):
    try:
        if bucket is None: return False
        blob_path = f"users/{user_email}/{file_name}"
        blob = bucket.blob(blob_path)
        blob.upload_from_string(file_bytes, content_type=file_type)
        if db:
            doc_ref = db.collection("users").document(user_email).collection("files").document()
            doc_ref.set({"name": file_name, "path": blob_path, "type": file_type, "created_at": firestore.SERVER_TIMESTAMP})
        return True
    except: return False

def download_from_firebase(file_path):
    try:
        if bucket is None: return None
        blob = bucket.blob(file_path)
        return blob.download_as_bytes()
    except: return None
    
# --- SOHBETGEÇMİŞİ (PERSISTENCE) ---
def create_new_research(user_email, title=None):
    """Yeni bir araştırma defteri oluşturur ve ID döner."""
    try:
        _title = title or ("Yeni Araştırma" if st.session_state.get("lang","TR") == "TR" else "New Research")
        ref = db.collection("users").document(user_email).collection("conversations").document()
        ref.set({
            "created_at": firestore.SERVER_TIMESTAMP,
            "title": _title,
            "sources": [],
        })
        return ref.id
    except: return None

# Geriye uyumluluk için alias
def create_new_conversation(user_email):
    return create_new_research(user_email, title="Yeni Sohbet")

def save_message_to_db(user_email, chat_id, role, content):
    """Mesajı veri tabanına kaydeder."""
    try:
        if not chat_id: return
        msg_ref = db.collection("users").document(user_email).collection("conversations").document(chat_id).collection("messages").document()
        msg_ref.set({
            "role": role,
            "content": content,
            "timestamp": firestore.SERVER_TIMESTAMP
        })
        
        # İlk mesajsa başlığı güncelle
        if role == "user":
            chat_doc = db.collection("users").document(user_email).collection("conversations").document(chat_id)
            # Sadece başlık "Yeni Sohbet" ise güncelle
            current_title = chat_doc.get().to_dict().get("title", "")
            if current_title == "Yeni Sohbet":
                new_title = content[:30] + "..." if len(content) > 30 else content
                chat_doc.update({"title": new_title})
                
    except Exception as e: print(f"Hata: {e}")

def get_user_conversations(user_email):
    """Araştırmaları listeler (yeniden eskiye)."""
    try:
        chats = db.collection("users").document(user_email).collection("conversations").order_by("created_at", direction=firestore.Query.DESCENDING).get()
        return [
            {
                "id": c.id,
                "title": c.to_dict().get("title", "İsimsiz"),
                "date": c.to_dict().get("created_at"),
                "sources": c.to_dict().get("sources", []),
            }
            for c in chats
        ]
    except: return []

def load_messages(user_email, chat_id):
    """Bir sohbete ait mesajları getirir."""
    try:
        msgs = db.collection("users").document(user_email).collection("conversations").document(chat_id).collection("messages").order_by("timestamp").get()
        return [{"role": m.to_dict()["role"], "content": m.to_dict()["content"]} for m in msgs]
    except: return []

def delete_conversation(user_email, chat_id):
    try:
        db.collection("users").document(user_email).collection("conversations").document(chat_id).delete()
        return True
    except: return False

def save_quiz_to_db(user_email, chat_id, quiz_data):
    """Sınav verisini kaydeder."""
    try:
        if not chat_id: return
        # Tek bir sınav belgesi olarak veya koleksiyon olarak tutabiliriz. 
        # Basitlik için sohbet dökümanının altına 'latest_quiz' olarak ekleyelim veya subcollection.
        # Subcollection daha esnek.
        batch = db.batch()
        quiz_ref = db.collection("users").document(user_email).collection("conversations").document(chat_id).collection("quizzes").document("latest")
        quiz_ref.set({"data": quiz_data, "timestamp": firestore.SERVER_TIMESTAMP})
    except: pass

def load_quiz_from_db(user_email, chat_id):
    try:
        doc = db.collection("users").document(user_email).collection("conversations").document(chat_id).collection("quizzes").document("latest").get()
        if doc.exists: return doc.to_dict().get("data", [])
        return []
    except: return []

def save_analysis_to_db(user_email, chat_id, analysis_data):
    """Analiz verisini kaydeder."""
    try:
        if not chat_id: return
        db.collection("users").document(user_email).collection("conversations").document(chat_id).collection("analyses").document("latest").set({
            "data": analysis_data, "timestamp": firestore.SERVER_TIMESTAMP
        })
    except: pass

def load_analysis_from_db(user_email, chat_id):
    try:
        doc = db.collection("users").document(user_email).collection("conversations").document(chat_id).collection("analyses").document("latest").get()
        if doc.exists: return doc.to_dict().get("data", None)
        return None
    except: return None

def save_map_to_db(user_email, chat_id, network_html, network_data):
    """Harita verisini kaydeder."""
    try:
        if not chat_id: return
        # HTML biraz büyük olabilir ama Firestore 1MB'a kadar izin veriyor.
        db.collection("users").document(user_email).collection("conversations").document(chat_id).collection("maps").document("latest").set({
            "html": network_html,
            "data": network_data,
            "timestamp": firestore.SERVER_TIMESTAMP
        })
    except Exception as e: print(f"Harita Kayıt Hatası: {e}")

def load_map_from_db(user_email, chat_id):
    try:
        doc = db.collection("users").document(user_email).collection("conversations").document(chat_id).collection("maps").document("latest").get()
        if doc.exists:
            d = doc.to_dict()
            return d.get("html", None), d.get("data", [])
        return None, []
    except: return None, []

def save_research_sources(user_email, research_id, docs):
    """Araştırmaya ait kaynak belgeleri Firestore'a kaydeder."""
    try:
        if not research_id: return
        db.collection("users").document(user_email).collection("conversations").document(research_id).update({
            "sources": docs
        })
    except: pass

def load_research_sources(user_email, research_id):
    """Araştırmadaki kaynak belgeleri yükler."""
    try:
        doc = db.collection("users").document(user_email).collection("conversations").document(research_id).get()
        if doc.exists: return doc.to_dict().get("sources", [])
        return []
    except: return []

def save_flashcards_to_db(user_email, research_id, cards):
    """Bilgi kartlarını Firestore'a kaydeder."""
    try:
        if not research_id: return
        db.collection("users").document(user_email).collection("conversations").document(research_id).collection("flashcards").document("latest").set({
            "data": cards, "timestamp": firestore.SERVER_TIMESTAMP
        })
    except: pass

def load_flashcards_from_db(user_email, research_id):
    """Bilgi kartlarını Firestore'dan yükler."""
    try:
        doc = db.collection("users").document(user_email).collection("conversations").document(research_id).collection("flashcards").document("latest").get()
        if doc.exists: return doc.to_dict().get("data", [])
        return []
    except: return []

def save_quiz_score_to_db(user_email, score, total, doc_names):
    """Quiz sonucunu genel skor geçmişine kaydeder."""
    try:
        db.collection("users").document(user_email).collection("quiz_history").add({
            "score": score,
            "total": total,
            "percentage": round((score / total) * 100) if total > 0 else 0,
            "documents": doc_names,
            "timestamp": firestore.SERVER_TIMESTAMP
        })
    except Exception as e: print(f"Skor Kayıt Hatası: {e}")

def get_quiz_history(user_email, limit=20):
    """Son N quiz sonucunu getirir (global)."""
    try:
        docs = (
            db.collection("users")
            .document(user_email)
            .collection("quiz_history")
            .order_by("timestamp", direction=firestore.Query.DESCENDING)
            .limit(limit)
            .get()
        )
        results = []
        for d in docs:
            data = d.to_dict()
            results.append({
                "score": data.get("score", 0),
                "total": data.get("total", 0),
                "percentage": data.get("percentage", 0),
                "documents": data.get("documents", []),
                "timestamp": data.get("timestamp"),
            })
        return list(reversed(results))
    except Exception as e:
        print(f"Skor Okuma Hatası: {e}")
        return []

def save_quiz_score_to_research(user_email, research_id, score, total, doc_names):
    """Quiz sonucunu araştırmaya özgü sub-collection'a kaydeder."""
    try:
        if not research_id: return
        db.collection("users").document(user_email).collection("conversations").document(research_id).collection("quiz_scores").add({
            "score": score,
            "total": total,
            "percentage": round((score / total) * 100) if total > 0 else 0,
            "documents": doc_names,
            "timestamp": firestore.SERVER_TIMESTAMP
        })
    except Exception as e: print(f"Araştırma Skor Kayıt: {e}")

def get_research_quiz_history(user_email, research_id, limit=20):
    """Bir araştırmaya ait quiz geçmişini getirir."""
    try:
        if not research_id: return []
        docs = (
            db.collection("users").document(user_email)
            .collection("conversations").document(research_id)
            .collection("quiz_scores")
            .order_by("timestamp", direction=firestore.Query.DESCENDING)
            .limit(limit).get()
        )
        results = []
        for d in docs:
            data = d.to_dict()
            results.append({
                "score":      data.get("score", 0),
                "total":      data.get("total", 0),
                "percentage": data.get("percentage", 0),
                "documents":  data.get("documents", []),
                "timestamp":  data.get("timestamp"),
            })
        return list(reversed(results))
    except Exception as e:
        print(f"Araştırma Skor Okuma: {e}")
        return []
        return []

SELECTED_MODEL = "gemini-flash-latest"

# --- YARDIMCI FONKSİYON: JSON TEMİZLEYİCİ ---
def clean_and_parse_json(text):
    try:
        text = text.replace("```json", "").replace("```", "").strip()
        start_idx = text.find('[')
        end_idx = text.rfind(']') + 1
        if start_idx != -1 and end_idx != 0:
            json_str = text[start_idx:end_idx]
            return json.loads(json_str)
        else:
            return []
    except json.JSONDecodeError:
        return [] 
    except: return []

# --- DIŞA AKTARMA FONKSİYONLARI ---
from fpdf import FPDF
from docx import Document
import io

def clean_text_for_pdf(text):
    """PDF için Türkçe karakterleri ASCII'ye çevirir (Basit çözüm)."""
    replacements = {
        "ğ": "g", "Ğ": "G", "ü": "u", "Ü": "U", "ş": "s", "Ş": "S",
        "ı": "i", "İ": "I", "ö": "o", "Ö": "O", "ç": "c", "Ç": "C"
    }
    for k, v in replacements.items():
        text = text.replace(k, v)
    return text.encode('latin-1', 'replace').decode('latin-1')

def create_pdf(text, title="Belge"):
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Arial", size=14)
        # Title
        pdf.cell(0, 10, txt=clean_text_for_pdf(title), ln=True, align='C')
        pdf.ln(10)
        # Body
        pdf.set_font("Arial", size=11)
        # Split text by newlines
        for line in text.split('\n'):
            # Latin-1 compatible check
            safe_line = clean_text_for_pdf(line)
            pdf.multi_cell(0, 7, txt=safe_line)
            
        return bytes(pdf.output())
    except Exception as e:
        print(f"PDF Hatası: {e}")
        return f"PDF Olusturulurken hata olustu: {e}".encode("utf-8")

def create_word(text, title="Belge"):

    doc = Document()
    doc.add_heading(title, 0)
    for line in text.split('\n'):
        doc.add_paragraph(line)
    
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def format_quiz_for_export(quiz_data):
    """Quiz verisini okunabilir metne dönüştürür."""
    text = ""
    for i, q in enumerate(quiz_data):
        text += f"{i+1}. {q['soru']}\n"
        for opt in q['secenekler']:
            text += f"   - {opt}\n"
        text += f"   (Cevap: {q['dogru_cevap']})\n\n"
    return text
 

def format_cards_for_export(cards):
    """Kartları okunabilir metne dönüştürür."""
    lines = []
    for i, c in enumerate(cards, 1):
        lines.append(f"{i}. {c.get('kavram','?')}  [{c.get('kategori','')}]")
        lines.append(f"   {c.get('tanim','?')}")
        lines.append("")
    return "\n".join(lines)

# --- AI & İŞLEME FONKSİYONLARI ---

def process_audio_data(audio_bytes, filename):
    try:
        with open("temp_audio.wav", "wb") as f: f.write(audio_bytes)
        myfile = genai.upload_file("temp_audio.wav")
        while myfile.state.name == "PROCESSING": time.sleep(1); myfile = genai.get_file(myfile.name)
        model = genai.GenerativeModel(SELECTED_MODEL)
        result = model.generate_content([myfile, "Ders notu çıkar."])
        text = result.text
        
        # RAG Kaydı
        if "rag_service" in st.session_state:
            success, msg = st.session_state.rag_service.ingest_text(text, source_name=filename)
            if success:
                st.session_state.rag_ready = True
                if filename not in st.session_state.selected_rag_docs:
                    st.session_state.selected_rag_docs.append(filename)
                st.success(msg)
            else:
                st.error(msg)
                
        if os.path.exists("temp_audio.wav"): os.remove("temp_audio.wav")
        return text
    except Exception as e: return f"Ses hatası: {str(e)}"

def process_pdf_data(file_bytes, filename):
    import io
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        total_text = ""
        for page in pdf_reader.pages:
            text = page.extract_text()
            if text: total_text += text + "\n"
        
        if len(total_text.strip()) < 10:
            st.error("⚠️ HATA: PDF okunamadı.")
            return ""

        # RAG Kaydı
        if "rag_service" in st.session_state:
            success, msg = st.session_state.rag_service.ingest_text(total_text, source_name=filename)
            if success:
                st.session_state.rag_ready = True
                if filename not in st.session_state.selected_rag_docs:
                    st.session_state.selected_rag_docs.append(filename)
                st.toast(msg, icon="✅")
            else:
                st.error(msg)
        
        return total_text
    except: return ""

def ask_gemini_rag(question, selected_docs=[]):
    try:
        if "rag_service" not in st.session_state:
            return "Hata: RAG servisi başlatılamadı."

        docs = st.session_state.rag_service.query(question, n_results=5, selected_sources=selected_docs)
        if not docs:
            return "🔍 Seçili belgelerde bu konuyla ilgili bilgi bulamadım."

        context_text = "\n\n".join([f"[Kaynak: {d.metadata.get('source', '')}] {d.page_content}" for d in docs])
        
        model = genai.GenerativeModel(SELECTED_MODEL)
        prompt = f"""
        Sen akademik bir asistansın. Aşağıdaki bağlama dayanarak soruya cevap ver.
        
        BAĞLAM (Veritabanından Bulunanlar):
        {context_text}
        
        SORU: {question}
        
        Cevabı Türkçe ver. Markdown formatını kullan.
        Cevabın içinde hangi kaynaktan bilgi aldığını belirtmek için [Kaynak Adı] formatını kullan.
        """
        response = model.generate_content(prompt)
        return response.text
    except Exception as e: return f"Hata: {e}"

def generate_quiz_from_docs(selected_docs=[]):
    try:
        if "rag_service" not in st.session_state: return []
        
        full_context = st.session_state.rag_service.get_full_text(selected_docs)
        if not full_context: return []

        model = genai.GenerativeModel(SELECTED_MODEL)
        prompt = f"""
        Aşağıdaki ders materyallerine dayanarak, zorluk derecesi yüksek 5 adet çoktan seçmeli sınav sorusu hazırla.
        Sorular metnin geneline yayılsın.
        
        MATERYAL (Başlangıç):
        {full_context[:500000]}
        
        Çıktı SADECE ve SADECE JSON formatında olsun.
        Format: [ {{"soru": "...", "secenekler": ["A) ...", "B) ..."], "dogru_cevap": "A) ..."}} ]
        """
        response = model.generate_content(prompt)
        return clean_and_parse_json(response.text)
    except Exception as e:
        st.error(f"Sınav Oluşturma Hatası: {e}")
        return []

def generate_flashcards_from_docs(selected_docs=[]):
    """Seçili belgelerden bilgi kartları (flashcard) üretir."""
    try:
        if "rag_service" not in st.session_state: return []
        full_context = st.session_state.rag_service.get_full_text(selected_docs)
        if not full_context: return []

        model = genai.GenerativeModel(SELECTED_MODEL)
        prompt = f"""
        Aşağıdaki ders materyalinden 15 adet bilgi kartı çıkar.
        Her kart üç alandan oluşsun:
          - "kavram": kısa terim veya başlık (max 8 kelime)
          - "tanim": net ve öğretici açıklama (2-4 cümle)
          - "kategori": tek kelimelik alan etiketi (ör. "Tanım", "Yöntem", "İlke", "Formül", "Tarih", "Şahsiyet")

        Kartlar belgenin farklı konularını kapsamalı.
        Çıktı SADECE ve SADECE JSON listesi olsun, başka hiçbir şey ekleme.
        Format:
        [
          {{"kavram": "...", "tanim": "...", "kategori": "..."}},
          ...
        ]

        MATERYAL:
        {full_context[:80000]}
        """
        response = model.generate_content(prompt)
        cards = clean_and_parse_json(response.text)
        return cards if isinstance(cards, list) else []
    except Exception as e:
        st.error(f"Kart Oluşturma Hatası: {e}")
        return []

def generate_network_from_docs(selected_docs=[]):
    try:
        if "rag_service" not in st.session_state: return
        
        full_context = st.session_state.rag_service.get_full_text(selected_docs)
        if not full_context: return

        model = genai.GenerativeModel(SELECTED_MODEL)
        prompt = f"""
        Aşağıdaki metindeki ana teknik kavramları ve aralarındaki ilişkileri JSON listesi olarak çıkar.
        En fazla 30 ilişki çıkar.
        
        MATERYAL:
        {full_context[:100000]}
        
        Çıktı SADECE JSON olsun.
        Format: [ {{"source": "Kavram A", "target": "Kavram B", "relation": "..."}} ]
        """
        response = model.generate_content(prompt)
        relationships = clean_and_parse_json(response.text)
        
        if not relationships:
            st.error("Harita verisi oluşturulamadı.")
            return
        
        # Datayı kaydet (Explorer için)
        st.session_state.network_data = relationships

        # --- NETWORK AYARLARI (GELİŞMİŞ) ---
        net = Network(height="600px", width="100%", bgcolor="#222222", font_color="white", notebook=False)
        options = {
            "nodes": {
                "font": {"size": 24, "color": "white", "face": "tahoma"}, 
                "shape": "dot",
                "size": 30 
            },
            "edges": {
                "width": 2,
                "color": {"color": "#aaaaaa", "highlight": "#ffffff"},
                "font": {"size": 18, "align": "middle", "color": "white", "strokeWidth": 2, "strokeColor": "black"} 
            },
            "physics": {
                "forceAtlas2Based": {
                    "gravitationalConstant": -100, 
                    "centralGravity": 0.005,      
                    "springLength": 200,          
                    "springConstant": 0.05,
                    "avoidOverlap": 1             
                },
                "minVelocity": 0.75,
                "solver": "forceAtlas2Based",
                "stabilization": {"enabled": True, "iterations": 1000} 
            },
            "interaction": {
                "navigationButtons": True, 
                "keyboard": True
            }
        }
        net.set_options(json.dumps(options))
        
        for item in relationships:
            net.add_node(item['source'], label=item['source'], color="#4ecdc4", title=item['source']) 
            net.add_node(item['target'], label=item['target'], color="#ff6b6b", title=item['target']) 
            net.add_edge(item['source'], item['target'], title=item['relation'], label=item['relation'])
            
        net.save_graph("network.html")
        with open("network.html", "r", encoding="utf-8") as f:
            st.session_state.network_html = f.read()
            
    except Exception as e:
        st.error(f"Harita Hatası: {e}")

def explain_relationship(node1, node2, selected_docs):
    try:
        full_context = st.session_state.rag_service.get_full_text(selected_docs)
        if not full_context: return "Metin bulunamadı."
        
        model = genai.GenerativeModel(SELECTED_MODEL)
        prompt = f"""
        Aşağıdaki metne dayanarak bu iki kavram arasındaki ilişkiyi detaylı açıkla.
        
        Kavram 1: {node1}
        Kavram 2: {node2}
        
        METİN:
        {full_context[:100000]}
        
        Cevabı Türkçe ver. Sadece gerçeği anlat.
        Markdown formatında, maddeler halinde yaz.
        """
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"Analiz hatası: {e}"

# --- LOGIN FUNC ---
def login_func(e, p):
    try:
        user = auth.sign_in_with_email_and_password(e, p)
        st.session_state.user = user
        st.session_state.rag_service = init_rag_for_user(user['email']) # RAG Yükle
        st.success("Giriş Başarılı!")
        time.sleep(0.5) 
        st.rerun()
    except Exception as err:
        st.error("Giriş başarısız. Lütfen bilgileri kontrol edin.")

def register_func(e, p):
    try:
        auth.create_user_with_email_and_password(e, p)
        st.success("Kayıt başarılı! Şimdi giriş yapabilirsiniz.")
    except: st.error("Kayıt hatası: Email kullanılıyor olabilir.")

def logout():
    st.session_state.user = None
    st.session_state.chat_history = []
    if "rag_service" in st.session_state: del st.session_state.rag_service
    st.session_state.rag_ready = False
    st.session_state.selected_rag_docs = []
    st.session_state.current_chat_id = None
    st.session_state.flashcards = []
    st.session_state.quiz_data = None
    st.session_state.network_html = None
    st.session_state.network_data = []
    st.rerun()

def _load_research(user_email, rid, sources):
    """Araştırmayı tam olarak session state'e yükler."""
    st.session_state.current_chat_id = rid
    st.session_state.selected_rag_docs = list(sources)
    # Checkbox state'lerini senkronize et ki otomatik seçim kalkmasın
    if "rag_service" in st.session_state and st.session_state.rag_service:
        for doc in st.session_state.rag_service.get_documents():
            st.session_state[f"chk_{doc}"] = (doc in sources)
            
    loaded_msgs = load_messages(user_email, rid)
    st.session_state.quiz_data = load_quiz_from_db(user_email, rid)
    st.session_state.notebook_explanation = load_analysis_from_db(user_email, rid)
    st.session_state.network_html, st.session_state.network_data = load_map_from_db(user_email, rid)
    st.session_state.flashcards = load_flashcards_from_db(user_email, rid)
    st.session_state.card_index = 0
    st.session_state.card_flipped = False
    st.session_state.chat_history = []
    for m in loaded_msgs:
        if m["role"] == "user": st.session_state.chat_history.append(HumanMessage(content=m["content"]))
        else: st.session_state.chat_history.append(AIMessage(content=m["content"]))

# --- MAIN APP ---
def main_app():
    user_email = st.session_state.user['email']
    short_email = user_email if len(user_email) < 28 else user_email[:25] + "..."

    # 1. TOP NAVBAR HİDERİ VE STİLİ
    st.markdown("""
        <style>
        .top-navbar-title { font-size: 1.8rem; font-weight: 800; color: #1E3A8A; margin:0; padding:0; display:flex; align-items:center; letter-spacing:-0.02em;}
        .top-navbar-sub { font-size: 0.85rem; color: #6B7280; margin-top: -5px; margin-bottom: 20px;}
        /* Hide the native sidebar completely for a clean look */
        [data-testid="stSidebar"] { display: none !important; }
        </style>
    """, unsafe_allow_html=True)
    
    active_count = len(st.session_state.selected_rag_docs)
    badge_html = (
        f'<span class="status-badge" style="font-size:0.75rem; background:#EFF6FF; color:#3B82F6; padding:3px 8px; border-radius:12px; margin-left:10px;">🟢 {t("badge_active", n=active_count)}</span>'
        if active_count > 0 else
        f'<span style="font-size:0.75rem; color:#EF4444; margin-left:10px;">{t("badge_none")}</span>'
    )

    st.markdown(f"<div class='top-navbar-title'>🎓 {t('app_title')} {badge_html}</div>", unsafe_allow_html=True)
    st.markdown(f"<div class='top-navbar-sub'>{t('app_subtitle')}</div>", unsafe_allow_html=True)

    # 2. ÜST MENÜ KOLONLARI (Açılır Menüler)
    col1, col2, col3 = st.columns([1.5, 1, 1])

    # ── KOLON 1: Araştırmalar & Belgeler (Popover) ──
    past_researches = get_user_conversations(user_email)
    
    with col1:
        with st.popover("📂 " + t('researches') + " & Belgeler", use_container_width=True):
            st.markdown(f"**{t('researches')}**")
            
            # Yeni Araştırma Ekle
            if st.button("➕ " + t("new_research"), key="btn_new_res", use_container_width=True, type="primary"):
                st.session_state["_show_research_form_pop"] = not st.session_state.get("_show_research_form_pop", False)
            if st.session_state.get("_show_research_form_pop", False):
                _default_title = t("untitled_research")
                _new_name = st.text_input(t("research_title_ph"), value="", placeholder=_default_title, key="_new_res_name_pop", label_visibility="collapsed")
                if st.button(t("new_research"), key="confirm_new_res_pop", use_container_width=True):
                    _title = _new_name.strip() or _default_title
                    new_rid = create_new_research(user_email, title=_title)
                    if new_rid:
                        st.session_state.current_chat_id = new_rid
                        st.session_state.chat_history = []
                        st.session_state.selected_rag_docs = []
                        st.session_state.active_page = "chat"
                        st.session_state["_show_research_form_pop"] = False
                        st.rerun()

            # Geçmiş Araştırmalar Listesi
            if not past_researches:
                st.caption(t("no_research_yet"))
            else:
                for r in past_researches:
                    rid = r['id']
                    is_active = st.session_state.current_chat_id == rid
                    r_title = r['title'] if len(r['title']) <= 25 else r['title'][:23] + "…"
                    type_btn = "primary" if is_active else "secondary"
                    # Sil ve Seç butonu yan yana
                    c_sel, c_del = st.columns([5, 1])
                    if c_sel.button(f"🔬 {r_title}", key=f"sel_r_pop_{rid}", use_container_width=True, type=type_btn):
                        _load_research(user_email, rid, r.get('sources', []))
                        st.session_state.active_page = "chat"
                        st.rerun()
                    if c_del.button("🗑", key=f"del_r_pop_{rid}"):
                        delete_conversation(user_email, rid)
                        if st.session_state.current_chat_id == rid:
                            st.session_state.current_chat_id = None
                        st.rerun()

            st.divider()

            # Kaynaklar (Mevcut Araştırmanın Seçili Belgeleri ve Tümü)
            available_docs = st.session_state.rag_service.get_documents()
            current_rid = st.session_state.current_chat_id
            st.markdown(f"**{t('research_sources')} ({len(st.session_state.selected_rag_docs)})**")
            
            if not available_docs:
                st.caption(t("no_doc_yet"))
            else:
                for doc in available_docs:
                    c1, c2 = st.columns([5, 1])
                    is_checked = doc in st.session_state.selected_rag_docs
                    lbl = doc if len(doc) < 22 else doc[:20] + "…"
                    new_val = c1.checkbox(lbl, value=is_checked, key=f"chk_pop_{doc}")
                    if new_val and not is_checked:
                        st.session_state.selected_rag_docs.append(doc)
                        if current_rid: save_research_sources(user_email, current_rid, st.session_state.selected_rag_docs)
                        st.rerun()
                    elif not new_val and is_checked:
                        st.session_state.selected_rag_docs.remove(doc)
                        if current_rid: save_research_sources(user_email, current_rid, st.session_state.selected_rag_docs)
                        st.rerun()
                    if c2.button("🗑", key=f"del2_{doc}"):
                        if st.session_state.rag_service.delete_document(doc):
                            if doc in st.session_state.selected_rag_docs:
                                st.session_state.selected_rag_docs.remove(doc)
                                if current_rid: save_research_sources(user_email, current_rid, st.session_state.selected_rag_docs)
                            st.toast(f"{doc} silindi.", icon="🗑️")
                            time.sleep(0.4)
                            st.rerun()
            
            st.divider()
            
            # Belge Yükle Section
            st.markdown(f"**{t('upload_doc')}**")
            source = st.radio(t("source_type"), [t("pdf_label"), t("audio_label")], horizontal=True, label_visibility="collapsed", key="popover_source")
            if source == t("pdf_label"):
                up_file = st.file_uploader(t("choose_pdf"), type="pdf", label_visibility="collapsed", key="pop_up_pdf")
                if up_file and st.button(t("upload_learn"), use_container_width=True, key="pop_btn_pdf"):
                    with st.spinner(t("reading")):
                        bd = up_file.read()
                        upload_to_firebase(bd, up_file.name, "application/pdf", user_email)
                        process_pdf_data(bd, up_file.name)
                        if up_file.name not in st.session_state.selected_rag_docs:
                            st.session_state.selected_rag_docs.append(up_file.name)
                            if current_rid: save_research_sources(user_email, current_rid, st.session_state.selected_rag_docs)
                        st.rerun()
            elif source == t("audio_label"):
                up_audio = st.file_uploader(t("choose_audio"), type=["mp3", "wav"], label_visibility="collapsed", key="pop_up_audio")
                if up_audio and st.button(t("process_learn"), use_container_width=True, key="pop_btn_audio"):
                    with st.spinner(t("processing")):
                        process_audio_data(up_audio.read(), up_audio.name)
                        upload_to_firebase(up_audio.getvalue(), up_audio.name, "audio/wav", user_email)
                        if up_audio.name not in st.session_state.selected_rag_docs:
                            st.session_state.selected_rag_docs.append(up_audio.name)
                            if current_rid: save_research_sources(user_email, current_rid, st.session_state.selected_rag_docs)
                        st.rerun()

    # ── KOLON 2: Sayfa / Özellik Seçici (Açılır Menü / Selectbox) ──
    with col2:
        _page_keys   = ["chat", "quiz", "map", "analytics", "cards"]
        _page_icons = {"chat": "💬", "quiz": "📝", "map": "🕸️", "analytics": "📊", "cards": "🃏"}
        _page_labels = {
            "chat":      t("nav_chat_title"),
            "quiz":      t("nav_quiz_title"),
            "map":       t("nav_map_title"),
            "analytics": t("nav_analytics_title"),
            "cards":     t("nav_cards_title"),
        }
        
        display_names = [ f"{_page_icons[k]} {_page_labels[k]}" for k in _page_keys ]
        current_idx = _page_keys.index(st.session_state.active_page) if st.session_state.active_page in _page_keys else 0
        
        selected_nav = st.selectbox(
            "Sayfa Seçin", 
            options=display_names, 
            index=current_idx, 
            label_visibility="collapsed"
        )
        new_page_index = display_names.index(selected_nav)
        new_page = _page_keys[new_page_index]
        if new_page != st.session_state.active_page:
            st.session_state.active_page = new_page
            st.rerun()

    # ── KOLON 3: Profil & Ayarlar (Popover) ──
    with col3:
        with st.popover(f"👤 {short_email}", use_container_width=True):
            # Dil Seçici
            lang_choice = st.radio(
                "", ["🌐 Türkçe", "🌍 English"],
                index=0 if st.session_state.lang == "TR" else 1,
                horizontal=True, label_visibility="collapsed", key="pop_lang"
            )
            new_lang = "TR" if "Türkçe" in lang_choice else "EN"
            if new_lang != st.session_state.lang:
                st.session_state.lang = new_lang
                st.rerun()
            
            st.divider()
            
            # Pomodoro Timer
            st.markdown(f"**🍅 {t('pomodoro_title')}**")
            pomodoro_html = f"""
            <div style="text-align: center; font-family: 'Inter', sans-serif; background: #FFFFFF; padding: 10px; border-radius: 8px; border: 1px solid #E5E7EB; box-shadow: 0 1px 2px rgba(0,0,0,0.02);">
                <div id="timer" style="font-size: 1.8rem; font-weight: 800; color: #3B82F6; margin: 5px 0; font-variant-numeric: tabular-nums; letter-spacing: -0.02em;">25:00</div>
                <div style="display: flex; justify-content: center; gap: 8px; margin-top: 10px;">
                    <button onclick="startTimer()" style="background: #3B82F6; color: white; border: none; padding: 6px 12px; border-radius: 6px; cursor: pointer; font-weight: 600; font-size: 0.85em; width: 100%;">{{t('pomodoro_start')}}</button>
                    <button onclick="resetTimer()" style="background: #F9FAFB; color: #4B5563; border: 1px solid #D1D5DB; padding: 6px 12px; border-radius: 6px; cursor: pointer; font-weight: 600; font-size: 0.85em; width: 100%;">{{t('pomodoro_reset')}}</button>
                </div>
                <script>
                    let timeLeft = 1500;
                    let timerId = null;
                    function updateDisplay() {{
                        let m = Math.floor(timeLeft / 60);
                        let s = timeLeft % 60;
                        document.getElementById('timer').innerText = (m < 10 ? "0" : "") + m + ":" + (s < 10 ? "0" : "") + s;
                    }}
                    function startTimer() {{
                        if(timerId !== null) return;
                        event.target.style.background = '#2563EB';
                        timerId = setInterval(() => {{
                            if(timeLeft > 0) {{
                                timeLeft--;
                                updateDisplay();
                            }} else {{
                                clearInterval(timerId);
                                timerId = null;
                                alert("{{t('pomodoro_break')}}");
                                timeLeft = 300; 
                                updateDisplay();
                            }}
                        }}, 1000);
                    }}
                    function resetTimer() {{
                        clearInterval(timerId);
                        timerId = null;
                        timeLeft = 1500;
                        updateDisplay();
                    }}
                </script>
            </div>
            """
            pomodoro_html_formatted = pomodoro_html.replace("{{t('pomodoro_title')}}", t('pomodoro_title')).replace("{{t('pomodoro_start')}}", t('pomodoro_start')).replace("{{t('pomodoro_reset')}}", t('pomodoro_reset')).replace("{{t('pomodoro_break')}}", t('pomodoro_break'))
            components.html(pomodoro_html_formatted, height=140)

            st.divider()

            if st.button(t("logout"), use_container_width=True, type="secondary"):
                logout()

    st.markdown("<hr style='border-top:1px solid #E5E7EB; margin: 15px 0 20px 0;'>", unsafe_allow_html=True)

    # ── MAVİ BİLGİ KARTUŞU (Active Research Banner) ──
    if st.session_state.current_chat_id:
        active_r = next((r for r in past_researches if r['id'] == st.session_state.current_chat_id), None)
        if active_r:
            a_title   = active_r['title']
            a_sources = active_r.get('sources', []) or st.session_state.selected_rag_docs
            chips_html = "".join(
                f'<span style="background:#F3F4F6;border:1px solid #E5E7EB;border-radius:20px;'
                f'padding:2px 9px;font-size:0.66rem;color:#4B5563;margin:2px 2px 0 0;display:inline-block;">'
                f'{s[:20]}</span>'
                for s in a_sources
            ) or f'<span style="font-size:0.7rem;color:#9CA3AF;">{{t("no_source_yet")}}</span>'
            st.markdown(
                f'<div style="background:#FFFFFF;border:1.5px solid #3B82F6;border-radius:10px;padding:10px 12px;margin-bottom:15px;box-shadow: 0 1px 3px rgba(0,0,0,0.05);">'
                f'<div style="font-size:0.78rem;font-weight:700;color:#1E3A8A;margin-bottom:5px">🔬 {a_title}</div>'
                f'<div>{chips_html}</div></div>',
                unsafe_allow_html=True
            )

    # ══════════════════════════════════════════
    # SAYFA: SOHBET
    # ══════════════════════════════════════════
    if st.session_state.active_page == "chat":
        # Geçmiş mesajları göster (belge seçili olmasa bile)
        if st.session_state.chat_history:
            for m in st.session_state.chat_history:
                role = "user" if isinstance(m, HumanMessage) else "assistant"
                avatar = "👤" if role == "user" else "🤖"
                with st.chat_message(role, avatar=avatar):
                    st.write(m.content)
        elif not available_docs:
            st.info(t("up_doc_first"))
        elif not st.session_state.selected_rag_docs:
            st.info(t("sel_doc_first"))

        # Yeni mesaj: sadece belge seçiliyse aktif
        if st.session_state.selected_rag_docs:
            if prompt := st.chat_input(t("chat_placeholder")):
                st.session_state.chat_history.append(HumanMessage(content=prompt))
                with st.chat_message("user", avatar="👤"):
                    st.write(prompt)
                with st.chat_message("assistant", avatar="🤖"):
                    with st.spinner(t("researching")):
                        response = ask_gemini_rag(prompt, selected_docs=st.session_state.selected_rag_docs)
                        st.write(response)
                        st.session_state.chat_history.append(AIMessage(content=response))
                        if st.session_state.current_chat_id is None:
                            st.session_state.current_chat_id = create_new_research(user_email)
                            if st.session_state.current_chat_id and st.session_state.selected_rag_docs:
                                save_research_sources(user_email, st.session_state.current_chat_id, st.session_state.selected_rag_docs)
                        if st.session_state.current_chat_id:
                            save_message_to_db(user_email, st.session_state.current_chat_id, "user", prompt)
                            save_message_to_db(user_email, st.session_state.current_chat_id, "assistant", response)
        else:
            st.chat_input(t("sel_doc_first"), disabled=True)

    # ══════════════════════════════════════════
    # SAYFA: SINAV
    # ══════════════════════════════════════════
    elif st.session_state.active_page == "quiz":
        st.subheader(t("quiz_title"))
        # Yeni sınav butonu — sadece belge seçiliyse aktif
        if st.session_state.selected_rag_docs:
            if st.button(t("quiz_create_btn", n=len(st.session_state.selected_rag_docs)), type="primary"):
                with st.spinner(t("quiz_spinner")):
                    st.session_state.quiz_data = generate_quiz_from_docs(st.session_state.selected_rag_docs)
                    st.session_state.quiz_answers = {}
                    st.session_state.quiz_submitted = False
                    if st.session_state.current_chat_id is None and st.session_state.quiz_data:
                        st.session_state.current_chat_id = create_new_conversation(user_email)
                    if st.session_state.current_chat_id:
                        save_quiz_to_db(user_email, st.session_state.current_chat_id, st.session_state.quiz_data)
        elif not st.session_state.quiz_data:
            st.info(t("quiz_sel_doc"))

        if st.session_state.quiz_data:
            total_q = len(st.session_state.quiz_data)

            # ── Sınav Tamamlandıysa Sonuç Göster ──
            if st.session_state.quiz_submitted:
                score = sum(1 for v in st.session_state.quiz_answers.values() if v)
                pct   = round((score / total_q) * 100)

                # Renk seç
                if pct >= 80:   color, emoji = "#2ecc71", "🏆"
                elif pct >= 50: color, emoji = "#f39c12", "📈"
                else:           color, emoji = "#e74c3c", "📉"

                st.markdown(f"""
                <div style="background:linear-gradient(135deg,#1a2332,#1e2a3a);border:2px solid {color};
                     border-radius:14px;padding:30px;text-align:center;margin:20px 0;">
                    <div style="font-size:3rem">{emoji}</div>
                    <div style="font-size:2.5rem;font-weight:800;color:{color}">{score}/{total_q}</div>
                    <div style="font-size:1.1rem;color:#aac4e8;margin-top:6px">{t('quiz_success_rate', pct=pct)}</div>
                </div>
                """, unsafe_allow_html=True)

                # Cevapları göster
                st.markdown("---")
                st.subheader(t("quiz_answer_key"))
                for i, q in enumerate(st.session_state.quiz_data):
                    is_correct = st.session_state.quiz_answers.get(i, False)
                    icon = "✅" if is_correct else "❌"
                    with st.expander(f"{icon} {i+1}. {q['soru'][:80]}…"):
                        for opt in q['secenekler']:
                            if opt == q['dogru_cevap']:
                                st.markdown(f"**✅ {opt}** {t('quiz_correct_ans')}")
                            else:
                                st.markdown(f"&nbsp;&nbsp;&nbsp;{opt}")

                col_r, col_n, _ = st.columns([1, 1, 2])
                if col_r.button(t("quiz_retry"), use_container_width=True):
                    st.session_state.quiz_answers = {}
                    st.session_state.quiz_submitted = False
                    st.rerun()
                if col_n.button("🔀 " + ("Yeni Sorular" if st.session_state.lang == "TR" else "New Questions"),
                                use_container_width=True, type="primary"):
                    if st.session_state.selected_rag_docs:
                        with st.spinner(t("quiz_spinner")):
                            st.session_state.quiz_data = generate_quiz_from_docs(st.session_state.selected_rag_docs)
                            st.session_state.quiz_answers = {}
                            st.session_state.quiz_submitted = False
                            if st.session_state.current_chat_id:
                                save_quiz_to_db(user_email, st.session_state.current_chat_id, st.session_state.quiz_data)
                            st.rerun()
                    else:
                        st.warning(t("quiz_sel_doc"))

            # ── Sınav Çözülüyor ──
            else:
                answered = len(st.session_state.quiz_answers)
                st.caption(t("quiz_progress", a=answered, t=total_q))
                st.progress(answered / total_q if total_q else 0)
                st.markdown("---")

                for i, q in enumerate(st.session_state.quiz_data):
                    is_answered = i in st.session_state.quiz_answers
                    st.markdown(f"**{i+1}. {q['soru']}**")
                    secim = st.radio(
                        "Seçiniz:", q['secenekler'],
                        key=f"rad_{i}", index=None,
                        label_visibility="collapsed"
                    )
                    if st.button(t("quiz_answer_btn", i=i+1), key=f"btn_{i}"):
                        if secim is None:
                            st.warning(t("quiz_choose_opt"))
                        else:
                            correct = (secim == q['dogru_cevap'])
                            st.session_state.quiz_answers[i] = correct
                            if correct: st.success(t("quiz_correct"))
                            else:       st.error(t("quiz_wrong", ans=q['dogru_cevap']))
                    st.divider()

                # Sınavı Bitir
                if st.button(t("quiz_finish"), type="primary", use_container_width=True):
                    if len(st.session_state.quiz_answers) < total_q:
                        st.warning(t("quiz_warn_unanswered", n=total_q - len(st.session_state.quiz_answers)))
                        if st.button(t("quiz_force_submit"), key="force_submit"):
                            st.session_state.quiz_submitted = True
                            score = sum(1 for v in st.session_state.quiz_answers.values() if v)
                            save_quiz_score_to_db(user_email, score, total_q, st.session_state.selected_rag_docs)
                            save_quiz_score_to_research(user_email, st.session_state.current_chat_id, score, total_q, st.session_state.selected_rag_docs)
                            st.rerun()
                    else:
                        st.session_state.quiz_submitted = True
                        score = sum(1 for v in st.session_state.quiz_answers.values() if v)
                        save_quiz_score_to_db(user_email, score, total_q, st.session_state.selected_rag_docs)
                        save_quiz_score_to_research(user_email, st.session_state.current_chat_id, score, total_q, st.session_state.selected_rag_docs)
                        st.rerun()


    # ══════════════════════════════════════════
    # SAYFA: KAVRAM HARİTASI
    # ══════════════════════════════════════════
    elif st.session_state.active_page == "map":
        st.subheader(t("map_title"))
        if st.session_state.selected_rag_docs:
            if st.button(t("map_draw"), type="primary"):
                with st.spinner(t("map_drawing")):
                    generate_network_from_docs(st.session_state.selected_rag_docs)
                    if st.session_state.current_chat_id is None and st.session_state.network_html:
                        st.session_state.current_chat_id = create_new_conversation(user_email)
                    if st.session_state.current_chat_id:
                        save_map_to_db(user_email, st.session_state.current_chat_id, st.session_state.network_html, st.session_state.network_data)
        elif not st.session_state.network_html:
            st.info(t("map_sel_doc"))

        if st.session_state.network_html:
            components.html(st.session_state.network_html, height=600)

            st.divider()
            st.subheader(t("map_analysis_title"))
            st.caption(t("map_analysis_desc"))

            if st.session_state.network_data:
                all_nodes = sorted(set(
                    [r['source'] for r in st.session_state.network_data] +
                    [r['target'] for r in st.session_state.network_data]
                ))
                c_sel1, c_sel2, c_act = st.columns([2, 2, 1])
                with c_sel1:
                    node1 = st.selectbox(t("node1_label"), all_nodes, key="node1")
                with c_sel2:
                    node2 = st.selectbox(t("node2_label"), [n for n in all_nodes if n != node1], key="node2")
                with c_act:
                    st.write(""); st.write("")
                    if st.button(t("analyze_btn")):
                        with st.spinner(t("analyzing")):
                            explanation = explain_relationship(node1, node2, st.session_state.selected_rag_docs)
                            st.session_state.notebook_explanation = {"n1": node1, "n2": node2, "text": explanation}
                            if st.session_state.current_chat_id is None:
                                st.session_state.current_chat_id = create_new_conversation(user_email)
                            if st.session_state.current_chat_id:
                                save_analysis_to_db(user_email, st.session_state.current_chat_id, st.session_state.notebook_explanation)

                if st.session_state.notebook_explanation:
                    exp = st.session_state.notebook_explanation
                    st.markdown(f"""
                    <div class="notebook-card">
                        <div class="notebook-header">
                            <span>🔗</span> {exp['n1']} ↔ {exp['n2']} Analizi
                        </div>
                        <div class="notebook-content">{exp['text']}</div>
                    </div>
                    """, unsafe_allow_html=True)

    # ══════════════════════════════════════════
    # SAYFA: QUIZ ANALİTİK
    # ══════════════════════════════════════════
    elif st.session_state.active_page == "analytics":
        st.subheader(t("analytics_title"))

        # Aktif araştırmaya özel geçmiş; araştırma seçili değilse global geçmiş
        _rid = st.session_state.current_chat_id
        history = get_research_quiz_history(user_email, _rid) if _rid else get_quiz_history(user_email, limit=20)

        if not history:
            st.info(t("analytics_no_data"))
        else:
            total_exams  = len(history)
            avg_pct      = round(sum(h["percentage"] for h in history) / total_exams)
            best_pct     = max(h["percentage"] for h in history)
            last_pct     = history[-1]["percentage"]

            # Metrik kartlar
            m1, m2, m3, m4 = st.columns(4)
            m1.metric(t("metric_total"), total_exams)
            m2.metric(t("metric_avg"),   f"%{avg_pct}" if st.session_state.lang == "TR" else f"{avg_pct}%")
            m3.metric(t("metric_best"),  f"%{best_pct}" if st.session_state.lang == "TR" else f"{best_pct}%")
            m4.metric(t("metric_last"),  f"%{last_pct}" if st.session_state.lang == "TR" else f"{last_pct}%")

            st.markdown("---")

            # Grafik (bar chart)
            st.subheader(t("trend_title"))
            import pandas as pd

            labels = [t("quiz_label", i=i+1) for i in range(len(history))]
            scores = [h["percentage"] for h in history]
            df_chart = pd.DataFrame({t("chart_col"): scores}, index=labels)
            st.bar_chart(df_chart, use_container_width=True, height=280)

            # Başarı bandı açıklaması
            col_a, col_b, col_c = st.columns(3)
            col_a.markdown(f'<div style="background:#2ecc7122;border-left:4px solid #2ecc71;padding:8px;border-radius:6px;font-size:0.8rem">{t("band_excellent")}</div>', unsafe_allow_html=True)
            col_b.markdown(f'<div style="background:#f39c1222;border-left:4px solid #f39c12;padding:8px;border-radius:6px;font-size:0.8rem">{t("band_improving")}</div>', unsafe_allow_html=True)
            col_c.markdown(f'<div style="background:#e74c3c22;border-left:4px solid #e74c3c;padding:8px;border-radius:6px;font-size:0.8rem">{t("band_retry")}</div>', unsafe_allow_html=True)

            st.markdown("---")

            # Detaylı tablo
            st.subheader(t("detail_title"))
            for i, h in enumerate(reversed(history)):
                pct = h["percentage"]
                if pct >= 80:   bar_color = "#2ecc71"
                elif pct >= 50: bar_color = "#f39c12"
                else:           bar_color = "#e74c3c"
                docs_str = ", ".join(h["documents"]) if h["documents"] else "—"
                ts = h["timestamp"]
                ts_str = ts.strftime("%d.%m.%Y %H:%M") if ts else "—"
                correct_str = t("detail_correct", s=h['score'], t=h['total'])

                st.markdown(f"""
                <div style="background:#1a202c;border:1px solid #2d3748;border-radius:10px;padding:14px 18px;margin-bottom:8px;display:flex;align-items:center;gap:16px;">
                    <div style="font-size:1.4rem;font-weight:800;color:{bar_color};min-width:60px">%{pct}</div>
                    <div style="flex:1">
                        <div style="font-size:0.85rem;color:#aac4e8;font-weight:600">{correct_str} &nbsp;·&nbsp; {ts_str}</div>
                        <div style="font-size:0.72rem;color:#556;margin-top:2px">📚 {docs_str[:80]}</div>
                    </div>
                    <div style="width:80px;background:#2d3748;border-radius:4px;height:8px;overflow:hidden">
                        <div style="width:{pct}%;background:{bar_color};height:100%;border-radius:4px"></div>
                    </div>
                </div>
                """, unsafe_allow_html=True)


    # ══════════════════════════════════════════
    # SAYFA: BİLGİ KARTLARI
    # ══════════════════════════════════════════
    elif st.session_state.active_page == "cards":
        st.subheader(t("cards_title"))

        if st.session_state.selected_rag_docs:
            col_gen, col_regen = st.columns([3, 1])
            with col_gen:
                if st.button(t("cards_create_btn", n=len(st.session_state.selected_rag_docs)), type="primary", use_container_width=True):
                    with st.spinner(t("cards_spinner")):
                        cards = generate_flashcards_from_docs(st.session_state.selected_rag_docs)
                        st.session_state.flashcards = cards
                        st.session_state.card_index = 0
                        st.session_state.card_flipped = False
                        if cards and st.session_state.current_chat_id:
                            save_flashcards_to_db(user_email, st.session_state.current_chat_id, cards)
            with col_regen:
                if st.session_state.flashcards and st.button(t("cards_regenerate"), use_container_width=True):
                    with st.spinner(t("cards_spinner")):
                        cards = generate_flashcards_from_docs(st.session_state.selected_rag_docs)
                        st.session_state.flashcards = cards
                        st.session_state.card_index = 0
                        st.session_state.card_flipped = False
                        if cards and st.session_state.current_chat_id:
                            save_flashcards_to_db(user_email, st.session_state.current_chat_id, cards)

        # Kartları göster — belge seçili olmasa da mevcut kartlar yüklüyse göster
        if not st.session_state.flashcards:
            if not st.session_state.selected_rag_docs:
                st.info(t("cards_sel_doc"))
            else:
                st.info("👆 Yukarıdaki butona basarak kartlarınızı oluşturun." if st.session_state.lang == "TR" else "👆 Press the button above to generate your cards.")
        else:
                cards    = st.session_state.flashcards
                idx      = st.session_state.card_index
                flipped  = st.session_state.card_flipped
                total    = len(cards)
                card     = cards[idx]
                kavram   = card.get("kavram", "—")
                tanim    = card.get("tanim",  "—")
                kategori = card.get("kategori", "")
                progress = round(((idx + 1) / total) * 100)

                # İlerleme çubuğu
                st.markdown(
                    f'<div class="fc-counter">{t("cards_of", cur=idx+1, total=total)}</div>'
                    f'<div class="fc-progress"><div class="fc-progress-bar" style="width:{progress}%"></div></div>',
                    unsafe_allow_html=True
                )

                # Flip card HTML
                flipped_class = "flipped" if flipped else ""
                kat_html  = f'<div class="fc-category">{kategori}</div>' if kategori else ""
                hint_text = "↩ Tanımı görmek için 'Kartı Çevir'e tıkla" if st.session_state.lang == "TR" else "↩ Click 'Flip Card' to see the definition"
                card_html = (
                    '<div class="fc-wrapper">'
                    f'  <div class="fc-inner {flipped_class}">'
                    '    <div class="fc-front">'
                    f'      {kat_html}'
                    f'      <div class="fc-term">{kavram}</div>'
                    f'      <div class="fc-hint">{hint_text}</div>'
                    '    </div>'
                    '    <div class="fc-back">'
                    f'      <div class="fc-definition">{tanim}</div>'
                    '    </div>'
                    '  </div>'
                    '</div>'
                    '<div style="min-height:310px"></div>'
                )
                st.markdown(card_html, unsafe_allow_html=True)

                # Kontrol butonları
                btn_prev, btn_flip, btn_next = st.columns([1, 2, 1])
                with btn_prev:
                    if st.button(t("cards_prev"), use_container_width=True, disabled=(idx == 0)):
                        st.session_state.card_index  -= 1
                        st.session_state.card_flipped = False
                        st.rerun()
                with btn_flip:
                    if st.button(t("cards_flip"), type="primary", use_container_width=True):
                        st.session_state.card_flipped = not flipped
                        st.rerun()
                with btn_next:
                    if st.button(t("cards_next"), use_container_width=True, disabled=(idx == total - 1)):
                        st.session_state.card_index  += 1
                        st.session_state.card_flipped = False
                        st.rerun()

                st.markdown("---")

                # Dışa aktarma
                ex_col1, ex_col2, _ = st.columns([1, 1, 2])
                export_text = format_cards_for_export(cards)
                with ex_col1:
                    pdf_bytes = create_pdf(export_text, title="Bilgi Kartları" if st.session_state.lang == "TR" else "Knowledge Cards")
                    st.download_button(t("cards_export_pdf"), data=pdf_bytes, file_name="bilgi_kartlari.pdf", mime="application/pdf", use_container_width=True)
                with ex_col2:
                    doc_bytes = create_word(export_text, title="Bilgi Kartları" if st.session_state.lang == "TR" else "Knowledge Cards")
                    st.download_button(t("cards_export_doc"), data=doc_bytes, file_name="bilgi_kartlari.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)

                # Tüm kartlar listesi
                with st.expander(t("cards_all")):
                    for i, c in enumerate(cards):
                        kat_badge = f'<span class="card-list-cat">{c.get("kategori","")}</span>' if c.get("kategori") else ""
                        st.markdown(
                            f'<div class="card-list-row">'
                            f'  <div class="card-list-num">{i+1}</div>'
                            f'  <div>'
                            f'    <div class="card-list-term">{c.get("kavram","")}</div>'
                            f'    {kat_badge}'
                            f'    <div class="card-list-def">{c.get("tanim","")}</div>'
                            f'  </div>'
                            f'</div>',
                            unsafe_allow_html=True
                        )


# --- LOGIN ---
if st.session_state.user is None:
    # Dil toggle (login ekranında da görünsün)
    _lc1, _lc2 = st.columns([2, 3])
    with _lc2:
        _lang_choice = st.radio(
            "", ["🌐 Türkçe", "🌍 English"],
            index=0 if st.session_state.lang == "TR" else 1,
            horizontal=True, label_visibility="collapsed"
        )
        _new_lang = "TR" if "Türkçe" in _lang_choice else "EN"
        if _new_lang != st.session_state.lang:
            st.session_state.lang = _new_lang
            st.rerun()

    c1, c2, c3 = st.columns([1,2,1])
    with c2:
        st.title(t("login_title"))
        tab1, tab2 = st.tabs([t("login_tab"), t("register_tab")])
        with tab1:
            e = st.text_input(t("email_label"), key="le")
            p = st.text_input(t("pass_label"), type="password", key="lp")
            if st.button(t("login_btn"), type="primary"): login_func(e, p)
        with tab2:
            e2 = st.text_input(t("email_label"), key="re")
            p2 = st.text_input(t("pass_label"), type="password", key="rp")
            if st.button(t("register_btn")): register_func(e2, p2)
else:
    main_app()